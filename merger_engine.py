"""
Document Merger Engine - Core merging logic
Handles PDF, DOCX, and Email merging with flexible folder structure support
"""

import atexit
import os
import json
from copy import deepcopy
from datetime import datetime, timezone
import shutil
import tempfile
import threading
import uuid
import zipfile
import traceback
from typing import List, Dict, Optional, Tuple, Set, Any, Callable
from collections import defaultdict
import re

# Module-level tracking of temp dirs for atexit cleanup if process is killed.
_active_temp_dirs: Set[str] = set()
_active_temp_dirs_lock = threading.Lock()


def _atexit_cleanup_temp_dirs():
    """Last-resort cleanup of temp dirs when the process exits."""
    with _active_temp_dirs_lock:
        for d in list(_active_temp_dirs):
            shutil.rmtree(d, ignore_errors=True)
        _active_temp_dirs.clear()


atexit.register(_atexit_cleanup_temp_dirs)

# PDF handling
try:
    from pypdf import PdfReader, PdfWriter
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

# Image handling (for converting images masquerading as PDFs)
try:
    from PIL import Image
    import io
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

# DOCX handling
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Email handling
try:
    import extract_msg
    HAS_EXTRACT_MSG = True
except ImportError:
    HAS_EXTRACT_MSG = False

# Microsoft Word automation (Windows)
try:
    import pythoncom
    import win32com.client as win32_client
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False

from email import policy
from email.parser import BytesParser

try:
    from dateutil import parser as date_parser
    HAS_DATEUTIL = True
except ImportError:
    date_parser = None  # type: ignore[assignment]
    HAS_DATEUTIL = False


def _record_warning(warnings: Optional[List[Dict]], code: str, message: str, **context) -> None:
    """Append a structured warning when a warning collector is provided."""
    if warnings is None:
        return
    warning = {'code': code, 'message': message}
    warning.update(context)
    warnings.append(warning)


def _safe_progress(callback, *args) -> None:
    """Call a progress callback, swallowing exceptions to avoid crashing the merge."""
    if callback is None:
        return
    try:
        callback(*args)
    except Exception:
        pass


def _make_writable_temp_dir(prefix: str) -> str:
    """
    Create a writable temporary directory.
    Some Windows/Python builds can produce temp dirs that are not writable when
    created with tempfile.mkdtemp(mode=0o700 semantics).
    """
    base_candidates = [tempfile.gettempdir(), os.getcwd()]

    for base_dir in base_candidates:
        if not base_dir:
            continue
        try:
            os.makedirs(base_dir, exist_ok=True)
        except OSError:
            continue

        for _ in range(8):
            candidate = os.path.join(base_dir, f"{prefix}{uuid.uuid4().hex}")
            try:
                os.makedirs(candidate, exist_ok=False)
                probe = os.path.join(candidate, ".write_probe")
                with open(probe, "wb") as handle:
                    handle.write(b"ok")
                os.remove(probe)
                return candidate
            except OSError:
                shutil.rmtree(candidate, ignore_errors=True)

    raise RuntimeError("Unable to create a writable temporary directory.")


class RunLogger:
    """Persist run events to text and JSONL logs."""

    def __init__(
        self,
        logs_dir: str,
        run_id: str,
        enabled: bool = True,
        privacy_mode: str = "redacted",
        event_callback: Optional[Callable[[Dict[str, Any]], None]] = None,
    ):
        self.enabled = enabled
        self.privacy_mode = privacy_mode
        self.run_id = run_id
        self.logs_dir = logs_dir
        self.event_callback = event_callback
        self.text_log_path = os.path.join(logs_dir, f"run_{run_id}.log")
        self.jsonl_log_path = os.path.join(logs_dir, f"run_{run_id}.jsonl")
        self._text_handle = None
        self._jsonl_handle = None

        if self.enabled:
            os.makedirs(self.logs_dir, exist_ok=True)
            self._text_handle = open(self.text_log_path, "a", encoding="utf-8")
            self._jsonl_handle = open(self.jsonl_log_path, "a", encoding="utf-8")

    def close(self) -> None:
        for handle in (self._text_handle, self._jsonl_handle):
            if handle is not None:
                try:
                    handle.close()
                except Exception:
                    pass

    def _redact_value(self, key: str, value):
        if self.privacy_mode != "redacted":
            return value
        if isinstance(value, str) and key.lower() in {"file", "source", "destination", "archive", "entry", "path"}:
            return os.path.basename(value)
        return value

    def _sanitize_context(self, context: Dict) -> Dict:
        sanitized = {}
        for key, value in context.items():
            sanitized[key] = self._redact_value(key, value)
        return sanitized

    def log(self, level: str, event: str, message: str, **context) -> None:
        if not self.enabled:
            return
        timestamp = datetime.now().isoformat()
        safe_context = self._sanitize_context(context)
        payload = {
            "ts": timestamp,
            "run_id": self.run_id,
            "level": level.upper(),
            "event": event,
            "message": message,
            "context": safe_context,
        }
        self._jsonl_handle.write(json.dumps(payload, ensure_ascii=False) + "\n")
        self._jsonl_handle.flush()

        text_context = ""
        if safe_context:
            context_parts = [f"{key}={value}" for key, value in sorted(safe_context.items())]
            text_context = " | " + ", ".join(context_parts)
        self._text_handle.write(f"[{timestamp}] {level.upper()} {event}: {message}{text_context}\n")
        self._text_handle.flush()
        if self.event_callback:
            try:
                self.event_callback(payload)
            except Exception:
                pass


class WordToPdfConverter:
    """Converts .doc/.docx files to PDF using Microsoft Word COM automation."""

    def __init__(self, warnings: Optional[List[Dict]] = None, timeout_seconds: int = 120):
        self.warnings = warnings
        self.word_app = None
        self.com_initialized = False
        self.timeout_seconds = timeout_seconds

    @staticmethod
    def is_available() -> Tuple[bool, str]:
        if os.name != 'nt':
            return False, "Word conversion is supported on Windows only."
        if not HAS_WIN32COM:
            return False, "pywin32 is required for Word-to-PDF conversion."
        return True, ""

    def __enter__(self):
        pythoncom.CoInitialize()
        self.com_initialized = True
        try:
            self.word_app = win32_client.DispatchEx("Word.Application")
            self.word_app.Visible = False
            self.word_app.DisplayAlerts = 0
            try:
                # 3 == msoAutomationSecurityForceDisable
                self.word_app.AutomationSecurity = 3
            except Exception:
                # If setting AutomationSecurity fails, continue with defaults.
                pass
            return self
        except Exception:
            # Ensure COM is uninitialized if initialization fails after CoInitialize.
            if self.word_app is not None:
                try:
                    self.word_app.Quit()
                except Exception:
                    pass
                finally:
                    self.word_app = None
            if self.com_initialized:
                pythoncom.CoUninitialize()
                self.com_initialized = False
            raise

    def __exit__(self, exc_type, exc, tb):
        if self.word_app is not None:
            try:
                self.word_app.Quit()
            except Exception:
                pass
            self.word_app = None
        if self.com_initialized:
            pythoncom.CoUninitialize()
            self.com_initialized = False
        return False

    def convert_file(self, source_path: str, output_pdf_path: str) -> bool:
        """Convert a single Word document to PDF. Returns True on success."""
        if self.word_app is None:
            raise RuntimeError("Word automation session is not initialized.")

        document = None
        timed_out = threading.Event()

        def _timeout_killer():
            """Force-close the document if conversion exceeds timeout."""
            timed_out.set()
            try:
                if document is not None:
                    document.Close(SaveChanges=False)
            except Exception:
                pass

        timer = threading.Timer(self.timeout_seconds, _timeout_killer)
        try:
            os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)
            input_abs = os.path.abspath(source_path)
            output_abs = os.path.abspath(output_pdf_path)

            timer.start()

            document = self.word_app.Documents.Open(
                input_abs,
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False,
                ConfirmConversions=False,
            )

            if timed_out.is_set():
                raise TimeoutError(f"Word conversion timed out after {self.timeout_seconds}s")

            # wdExportFormatPDF = 17
            if hasattr(document, "ExportAsFixedFormat"):
                document.ExportAsFixedFormat(output_abs, 17)
            else:
                # Fallback for older Office object models.
                document.SaveAs(output_abs, FileFormat=17)

            if timed_out.is_set():
                raise TimeoutError(f"Word conversion timed out after {self.timeout_seconds}s")

            return os.path.exists(output_abs) and os.path.getsize(output_abs) > 0
        except TimeoutError as exc:
            _record_warning(
                self.warnings,
                'word_to_pdf_timeout',
                f'Word conversion timed out after {self.timeout_seconds}s; skipping file',
                file=source_path,
                timeout=self.timeout_seconds,
            )
            return False
        except Exception as exc:
            _record_warning(
                self.warnings,
                'word_to_pdf_failed',
                'Word-to-PDF conversion failed; skipping file',
                file=source_path,
                error=str(exc),
            )
            return False
        finally:
            timer.cancel()
            if document is not None:
                try:
                    document.Close(SaveChanges=False)
                except Exception:
                    pass


class PDFMerger:
    """Merges multiple PDF files into batched output files"""
    
    def __init__(self, max_file_size_kb=102400):
        self.max_file_size_kb = max_file_size_kb
        self.max_file_size_bytes = max_file_size_kb * 1024
        
    def estimate_batch_count(self, pdf_files: List[str]) -> int:
        """Estimate how many output batches a merge operation will create."""
        pdf_files = sorted(pdf_files)
        if not pdf_files:
            return 0

        batches = 1
        current_batch_size = 0

        for pdf_file in pdf_files:
            try:
                file_size = os.path.getsize(pdf_file)
            except OSError:
                # Avoid underestimating in preflight checks.
                file_size = self.max_file_size_bytes

            if current_batch_size and (current_batch_size + file_size > self.max_file_size_bytes):
                batches += 1
                current_batch_size = 0

            current_batch_size += file_size

        return batches

    def merge_pdfs(
        self,
        pdf_files: List[str],
        output_path: str,
        group_name: str,
        warnings: Optional[List[Dict]] = None,
        output_label: str = "pdfs",
        bookmark_titles: Optional[Dict[str, str]] = None,
        source_file_map: Optional[Dict[str, str]] = None,
        output_to_sources: Optional[Dict[str, List[str]]] = None,
    ) -> List[str]:
        """
        Merge PDF files into batches, staying under size limit
        
        Args:
            pdf_files: List of PDF file paths to merge
            output_path: Directory to save merged PDFs
            group_name: Name prefix for output files (e.g., "case_12345")
            
        Returns:
            List of created output file paths
        """
        if not HAS_PYPDF:
            raise ImportError("pypdf library is required for PDF merging")
        
        os.makedirs(output_path, exist_ok=True)
        output_files = []

        # Sort PDFs by name for consistent ordering
        pdf_files = sorted(pdf_files)

        current_batch = []
        current_batch_size = 0
        current_batch_words = 0
        max_batch_words = 50000  # netdoc word limit
        batch_num = 1

        for pdf_file in pdf_files:
            try:
                file_size = os.path.getsize(pdf_file)
            except OSError as exc:
                _record_warning(
                    warnings,
                    'pdf_stat_failed',
                    'Could not determine PDF file size; using max batch size for pre-allocation',
                    file=pdf_file,
                    error=str(exc),
                )
                file_size = self.max_file_size_bytes

            # Extract text to count words
            file_text = self._extract_pdf_text(pdf_file)
            file_words = self._count_words_in_text(file_text)

            # If adding this file would exceed limit, save current batch
            if current_batch and (current_batch_size + file_size > self.max_file_size_bytes or
                                  current_batch_words + file_words > max_batch_words):
                output_file = self._save_pdf_batch(
                    current_batch,
                    output_path,
                    group_name,
                    batch_num,
                    warnings,
                    output_label=output_label,
                    bookmark_titles=bookmark_titles,
                    source_file_map=source_file_map,
                    output_to_sources=output_to_sources,
                )
                if output_file:
                    output_files.append(output_file)
                batch_num += 1
                current_batch = []
                current_batch_size = 0
                current_batch_words = 0

            current_batch.append(pdf_file)
            current_batch_size += file_size
            current_batch_words += file_words

            # Warn if single file exceeds word limit
            if file_words > max_batch_words:
                _record_warning(
                    warnings,
                    'pdf_exceeds_word_cap',
                    'PDF exceeds netdoc word limit (50000); writing dedicated batch file',
                    file=pdf_file,
                    file_words=file_words,
                    batch_limit_words=max_batch_words,
                )
        
        # Save remaining files
        if current_batch:
            output_file = self._save_pdf_batch(
                current_batch,
                output_path,
                group_name,
                batch_num,
                warnings,
                output_label=output_label,
                bookmark_titles=bookmark_titles,
                source_file_map=source_file_map,
                output_to_sources=output_to_sources,
            )
            if output_file:
                output_files.append(output_file)
        
        return output_files
    
    def _try_convert_image_to_pdf(self, file_path: str) -> Optional[bytes]:
        """
        Attempt to open a file as an image and convert it to PDF bytes.
        Returns PDF bytes on success, or None if the file is not a valid image.
        """
        if not HAS_PIL:
            return None
        try:
            img = Image.open(file_path)
            # Convert to RGB so it can be saved as PDF (handles RGBA, P, etc.)
            if img.mode not in ('RGB', 'L'):
                img = img.convert('RGB')
            pdf_bytes = io.BytesIO()
            img.save(pdf_bytes, format='PDF', resolution=150)
            print(f"    Converted image to PDF: {os.path.basename(file_path)}")
            return pdf_bytes.getvalue()
        except Exception:
            return None

    def _try_convert_ole_doc_to_pdf(self, file_path: str) -> Optional[bytes]:
        """
        Attempt to extract text from an OLE (legacy .doc) file and render it
        as a simple PDF page. Returns PDF bytes on success, or None on failure.
        """
        if not HAS_PIL:
            return None
        try:
            import olefile
        except ImportError:
            return None

        try:
            ole = olefile.OleFileIO(file_path)
            # Extract text from the WordDocument stream
            text = ""
            if ole.exists('WordDocument'):
                # Try to get text from the Word Document stream
                # The actual text in .doc files is in a complex binary format,
                # but we can try extracting readable ASCII/Unicode content
                raw = ole.openstream('WordDocument').read()
                # Extract printable text chunks
                text_chunks = []
                current_chunk = []
                for byte in raw:
                    if 32 <= byte < 127 or byte in (10, 13, 9):
                        current_chunk.append(chr(byte))
                    else:
                        if len(current_chunk) > 3:  # Only keep chunks > 3 chars
                            text_chunks.append(''.join(current_chunk))
                        current_chunk = []
                if len(current_chunk) > 3:
                    text_chunks.append(''.join(current_chunk))
                text = '\n'.join(text_chunks)
            ole.close()

            if not text.strip():
                return None

            # Render text as an image then convert to PDF
            from PIL import ImageDraw, ImageFont
            # Create a page-sized image (A4 at 72dpi equivalent)
            page_width, page_height = 612, 792
            margin = 40
            line_height = 14
            max_chars_per_line = 80
            max_lines = (page_height - 2 * margin) // line_height

            # Word-wrap the text
            import textwrap
            lines = []
            for paragraph in text.split('\n'):
                wrapped = textwrap.wrap(paragraph, width=max_chars_per_line) or ['']
                lines.extend(wrapped)

            # Create pages
            pages = []
            for page_start in range(0, len(lines), max_lines):
                page_lines = lines[page_start:page_start + max_lines]
                img = Image.new('RGB', (page_width, page_height), 'white')
                draw = ImageDraw.Draw(img)
                # Add header
                basename = os.path.basename(file_path)
                draw.text((margin, margin // 2), f"[Extracted from: {basename}]", fill='gray')
                for i, line in enumerate(page_lines):
                    y = margin + i * line_height
                    draw.text((margin, y), line, fill='black')
                pages.append(img)

            if not pages:
                return None

            # Save all pages as a multi-page PDF
            pdf_bytes = io.BytesIO()
            if len(pages) == 1:
                pages[0].save(pdf_bytes, format='PDF', resolution=72)
            else:
                pages[0].save(pdf_bytes, format='PDF', resolution=72,
                              save_all=True, append_images=pages[1:])
            print(f"    Converted OLE doc to PDF: {os.path.basename(file_path)}")
            return pdf_bytes.getvalue()
        except Exception:
            return None

    def _extract_pdf_text(self, pdf_file: str) -> str:
        """Extract text from a PDF file for word counting."""
        try:
            reader = PdfReader(pdf_file)
            if reader.is_encrypted:
                # Try to decrypt with empty password (handles "view-only" PDFs)
                result = reader.decrypt("")
                if not result:
                    return ""  # genuinely password-protected, can't read
            text = ""
            for page in reader.pages:
                try:
                    text += page.extract_text() or ""
                except Exception:
                    pass
            return text
        except Exception:
            return ""

    def _count_words_in_text(self, text: str) -> int:
        """Count words in text by splitting on whitespace."""
        return len(text.split())

    def _save_pdf_batch(
        self,
        pdf_files: List[str],
        output_path: str,
        group_name: str,
        batch_num: int,
        warnings: Optional[List[Dict]] = None,
        output_label: str = "pdfs",
        bookmark_titles: Optional[Dict[str, str]] = None,
        source_file_map: Optional[Dict[str, str]] = None,
        output_to_sources: Optional[Dict[str, List[str]]] = None,
    ) -> Optional[str]:
        """Save a batch of PDFs into a single merged PDF"""
        writer = PdfWriter()
        total_pages_added = 0
        merged_batch_sources: List[str] = []

        def add_bookmark(title: str, page_index: int) -> None:
            if page_index < 0:
                return
            try:
                writer.add_outline_item(title, page_index)
            except Exception:
                try:
                    writer.addBookmark(title, page_index)
                except Exception:
                    pass
        
        # Add all pages from all PDFs
        for pdf_file in pdf_files:
            try:
                reader = PdfReader(pdf_file)
                if reader.is_encrypted:
                    # Try to decrypt with empty password (handles "view-only" PDFs)
                    result = reader.decrypt("")
                    if not result:
                        # genuinely password-protected, can't decrypt
                        _record_warning(
                            warnings,
                            'pdf_encrypted',
                            'PDF is password-protected and cannot be merged; skipping',
                            file=pdf_file,
                        )
                        continue
                    # else: successfully decrypted, continue with normal processing
                page_start = total_pages_added
                file_pages_added = 0
                for page in reader.pages:
                    writer.add_page(page)
                    file_pages_added += 1
                if file_pages_added == 0:
                    _record_warning(
                        warnings,
                        'pdf_no_pages',
                        'PDF contained zero readable pages',
                        file=pdf_file,
                    )
                else:
                    if bookmark_titles and bookmark_titles.get(pdf_file):
                        add_bookmark(bookmark_titles[pdf_file], page_start)
                    merged_batch_sources.append(pdf_file)
                total_pages_added += file_pages_added
            except Exception as e:
                # Fallback 1: File may be an image with a .pdf extension
                pdf_bytes = self._try_convert_image_to_pdf(pdf_file)
                if not pdf_bytes:
                    # Fallback 2: File may be an OLE .doc with a .pdf extension
                    pdf_bytes = self._try_convert_ole_doc_to_pdf(pdf_file)
                if pdf_bytes:
                    try:
                        reader = PdfReader(io.BytesIO(pdf_bytes))
                        page_start = total_pages_added
                        converted_pages = 0
                        for page in reader.pages:
                            writer.add_page(page)
                            converted_pages += 1
                        total_pages_added += converted_pages
                        if converted_pages == 0:
                            _record_warning(
                                warnings,
                                'pdf_conversion_empty',
                                'Fallback conversion produced zero pages',
                                file=pdf_file,
                            )
                        else:
                            if bookmark_titles and bookmark_titles.get(pdf_file):
                                add_bookmark(bookmark_titles[pdf_file], page_start)
                            merged_batch_sources.append(pdf_file)
                    except Exception as e2:
                        _record_warning(
                            warnings,
                            'pdf_conversion_failed',
                            'Could not merge file after fallback conversion',
                            file=pdf_file,
                            error=str(e2),
                        )
                        print(f"Warning: Could not merge {pdf_file} even after conversion: {e2}")
                else:
                    _record_warning(
                        warnings,
                        'pdf_unreadable',
                        'Could not read PDF file and fallback conversion failed',
                        file=pdf_file,
                        error=str(e),
                    )
                    print(f"Warning: Could not merge {pdf_file}: {e}")

        if total_pages_added == 0:
            _record_warning(
                warnings,
                'pdf_empty_batch',
                'Skipped PDF batch because no readable pages were found',
                group=group_name,
                batch=batch_num,
                file_count=len(pdf_files),
            )
            print(f"Warning: Skipping empty PDF batch {batch_num} for group {group_name}")
            return None
        
        # Generate output filename
        output_filename = f"{group_name}_{output_label}_batch{batch_num}.pdf"
        output_file = os.path.join(output_path, output_filename)

        if output_to_sources is not None:
            mapped_sources: List[str] = []
            for merged_source in merged_batch_sources:
                original = source_file_map.get(merged_source, merged_source) if source_file_map else merged_source
                if original not in mapped_sources:
                    mapped_sources.append(original)
            output_to_sources[output_file] = mapped_sources
        
        # Write merged PDF
        with open(output_file, 'wb') as f:
            writer.write(f)
        
        print(f"    Created: {output_filename} ({len(pdf_files)} PDFs, {total_pages_added} pages)")
        return output_file


class DOCXMerger:
    """Merges multiple DOCX files into batched output files"""
    
    def __init__(self, max_file_size_kb=102400):
        self.max_file_size_kb = max_file_size_kb
        self.max_file_size_bytes = max_file_size_kb * 1024
        
    def estimate_batch_count(self, docx_files: List[str]) -> int:
        """Estimate how many output batches a DOCX merge operation will create."""
        docx_files = sorted(docx_files)
        if not docx_files:
            return 0

        batches = 1
        current_batch_size = 0

        for docx_file in docx_files:
            try:
                file_size = os.path.getsize(docx_file)
            except OSError:
                file_size = self.max_file_size_bytes

            if current_batch_size and (current_batch_size + file_size > self.max_file_size_bytes):
                batches += 1
                current_batch_size = 0

            current_batch_size += file_size

        return batches

    def merge_docx(
        self,
        docx_files: List[str],
        output_path: str,
        group_name: str,
        warnings: Optional[List[Dict]] = None,
    ) -> List[str]:
        """
        Merge DOCX files into batches, staying under size limit
        
        Args:
            docx_files: List of DOCX file paths to merge
            output_path: Directory to save merged DOCX files
            group_name: Name prefix for output files
            
        Returns:
            List of created output file paths
        """
        if not HAS_DOCX:
            raise ImportError("python-docx library is required for DOCX merging")
        
        os.makedirs(output_path, exist_ok=True)
        output_files = []

        # Sort files by name
        docx_files = sorted(docx_files)

        current_batch = []
        current_batch_size = 0
        current_batch_words = 0
        max_batch_words = 50000  # netdoc word limit
        batch_num = 1

        for docx_file in docx_files:
            try:
                file_size = os.path.getsize(docx_file)
            except OSError as exc:
                _record_warning(
                    warnings,
                    'docx_stat_failed',
                    'Could not determine document file size; using max batch size for pre-allocation',
                    file=docx_file,
                    error=str(exc),
                )
                file_size = self.max_file_size_bytes

            # Extract text to count words
            file_text = self._extract_docx_text_for_counting(docx_file)
            file_words = self._count_words_in_text(file_text)

            # Check if we need to start a new batch
            if current_batch and (current_batch_size + file_size > self.max_file_size_bytes or
                                  current_batch_words + file_words > max_batch_words):
                output_file = self._save_docx_batch(
                    current_batch, output_path, group_name, batch_num, warnings
                )
                if output_file:
                    output_files.append(output_file)
                batch_num += 1
                current_batch = []
                current_batch_size = 0
                current_batch_words = 0

            current_batch.append(docx_file)
            current_batch_size += file_size
            current_batch_words += file_words

            # Warn if single file exceeds word limit
            if file_words > max_batch_words:
                _record_warning(
                    warnings,
                    'docx_exceeds_word_cap',
                    'DOCX exceeds netdoc word limit (50000); writing dedicated batch file',
                    file=docx_file,
                    file_words=file_words,
                    batch_limit_words=max_batch_words,
                )
        
        # Save remaining files
        if current_batch:
            output_file = self._save_docx_batch(
                current_batch, output_path, group_name, batch_num, warnings
            )
            if output_file:
                output_files.append(output_file)
        
        return output_files
    
    def _try_extract_docx_text(self, file_path: str) -> Optional[str]:
        """
        Fallback: extract raw paragraph text from word/document.xml inside the zip.
        Works even when python-docx can't open the file due to broken relationships.
        """
        import zipfile
        import xml.etree.ElementTree as ET

        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                # Find word/document.xml (may be at a different path in some files)
                candidates = [n for n in z.namelist() if n.endswith('document.xml')]
                if not candidates:
                    return None
                xml_bytes = z.read(candidates[0])
            root = ET.fromstring(xml_bytes)
            paragraphs = []
            for para in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                texts = [t.text or '' for t in para.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')]
                line = ''.join(texts)
                paragraphs.append(line)
            return '\n'.join(paragraphs) if paragraphs else None
        except Exception:
            return None

    def _try_extract_ole_text(self, file_path: str) -> Optional[str]:
        """
        Fallback: extract text from OLE (legacy .doc) compound files.
        These may appear as OOXML theme-only zips but contain actual content
        in OLE streams like WordDocument and 1Table.
        """
        try:
            import olefile
        except ImportError:
            return None
        try:
            if not olefile.isOleFile(file_path):
                return None
            ole = olefile.OleFileIO(file_path)
            text = ""
            if ole.exists('WordDocument'):
                raw = ole.openstream('WordDocument').read()
                text_chunks = []
                current_chunk = []
                for byte in raw:
                    if 32 <= byte < 127 or byte in (10, 13, 9):
                        current_chunk.append(chr(byte))
                    else:
                        if len(current_chunk) > 3:
                            text_chunks.append(''.join(current_chunk))
                        current_chunk = []
                if len(current_chunk) > 3:
                    text_chunks.append(''.join(current_chunk))
                text = '\n'.join(text_chunks)
            ole.close()
            return text.strip() if text.strip() else None
        except Exception:
            return None

    def _extract_docx_text_for_counting(self, file_path: str) -> str:
        """Extract text from DOCX file for word counting."""
        # Try using python-docx library
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + " "
            return text
        except Exception:
            pass

        # Fallback to raw text extraction
        raw_text = self._try_extract_docx_text(file_path)
        if raw_text:
            return raw_text
        # Try OLE fallback
        ole_text = self._try_extract_ole_text(file_path)
        if ole_text:
            return ole_text
        return ""

    def _count_words_in_text(self, text: str) -> int:
        """Count words in text by splitting on whitespace."""
        return len(text.split())

    def _save_docx_batch(
        self,
        docx_files: List[str],
        output_path: str,
        group_name: str,
        batch_num: int,
        warnings: Optional[List[Dict]] = None,
    ) -> Optional[str]:
        """Save a batch of DOCX files into a single merged document."""
        import tempfile
        import shutil

        merged_doc = Document()
        merged_docs_count = 0

        for docx_file in docx_files:
            source_doc = None
            raw_text = None
            open_error = None

            # Attempt 1: open normally.
            try:
                source_doc = Document(docx_file)
            except Exception as e1:
                open_error = e1

                # Attempt 2: copy to a temp .docx and retry.
                # Handles .doc files that are actually OOXML with wrong extension.
                tmp_path = None
                try:
                    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
                    os.close(tmp_fd)
                    shutil.copy2(docx_file, tmp_path)
                    source_doc = Document(tmp_path)
                    print(f"    Recovered (as .docx): {os.path.basename(docx_file)}")
                except Exception:
                    source_doc = None
                finally:
                    if tmp_path and os.path.exists(tmp_path):
                        try:
                            os.remove(tmp_path)
                        except Exception:
                            pass

                if source_doc is None:
                    # Attempt 3: extract raw text from the zip's document.xml.
                    raw_text = self._try_extract_docx_text(docx_file)
                    if raw_text:
                        print(f"    Recovered (raw text): {os.path.basename(docx_file)}")
                    else:
                        # Attempt 4: extract text from OLE compound document.
                        raw_text = self._try_extract_ole_text(docx_file)
                        if raw_text:
                            print(f"    Recovered (OLE text): {os.path.basename(docx_file)}")

            if source_doc is None and not raw_text:
                _record_warning(
                    warnings,
                    'docx_unreadable',
                    'Could not read DOCX/DOC file; skipping file',
                    file=docx_file,
                    error=str(open_error) if open_error else 'unknown_error',
                )
                print(f"Warning: Could not merge {docx_file}: {open_error}")
                continue

            if raw_text:
                if merged_docs_count > 0:
                    merged_doc.add_page_break()
                heading = merged_doc.add_heading(level=1)
                heading.text = f"Document: {os.path.basename(docx_file)}"
                merged_doc.add_paragraph(raw_text)
                merged_docs_count += 1
                continue

            try:
                elements = [
                    deepcopy(element)
                    for element in source_doc.element.body.iterchildren()
                    if not element.tag.endswith('}sectPr')
                ]
                if not elements:
                    _record_warning(
                        warnings,
                        'docx_empty_document',
                        'DOCX file had no readable body elements; skipping file',
                        file=docx_file,
                    )
                    continue

                if merged_docs_count > 0:
                    merged_doc.add_page_break()
                heading = merged_doc.add_heading(level=1)
                heading.text = f"Document: {os.path.basename(docx_file)}"

                for element in elements:
                    merged_doc.element.body.append(element)
                merged_docs_count += 1
            except Exception as e:
                _record_warning(
                    warnings,
                    'docx_append_failed',
                    'Could not append DOCX body elements; skipping file',
                    file=docx_file,
                    error=str(e),
                )
                print(f"Warning: Could not append content from {docx_file}: {e}")

        if merged_docs_count == 0:
            _record_warning(
                warnings,
                'docx_empty_batch',
                'Skipped DOCX batch because no readable documents were found',
                group=group_name,
                batch=batch_num,
                file_count=len(docx_files),
            )
            print(f"Warning: Skipping empty DOCX batch {batch_num} for group {group_name}")
            return None

        # Generate output filename
        output_filename = f"{group_name}_documents_batch{batch_num}.docx"
        output_file = os.path.join(output_path, output_filename)

        # Save merged document
        merged_doc.save(output_file)

        print(f"    Created: {output_filename} ({merged_docs_count} documents)")
        return output_file


class EmailExtractor:
    """Extracts email content and metadata from .msg and .eml files"""
    
    @staticmethod
    def extract_msg(file_path: str) -> Optional[Dict]:
        """Extract email data from .msg file"""
        if not HAS_EXTRACT_MSG:
            return None
            
        msg = None
        try:
            msg = extract_msg.Message(file_path)
            attachments = []
            for attachment in getattr(msg, "attachments", []) or []:
                filename = (
                    getattr(attachment, "longFilename", None)
                    or getattr(attachment, "shortFilename", None)
                    or "unnamed_attachment"
                )
                content_type = getattr(attachment, "mimeType", None) or ""
                size_bytes = None
                data = getattr(attachment, "data", None)
                if data is not None:
                    try:
                        size_bytes = len(data)
                    except Exception:
                        size_bytes = None
                attachments.append(
                    {
                        "filename": filename,
                        "content_type": content_type,
                        "size_bytes": size_bytes,
                    }
                )
            return {
                'subject': msg.subject or '(No Subject)',
                'from': msg.sender or '',
                'to': msg.to or '',
                'cc': msg.cc or '',
                'date': msg.date,
                'body': msg.body or '',
                'attachments': attachments,
            }
        except Exception as e:
            print(f"Error extracting .msg file {file_path}: {e}")
            return None
        finally:
            if msg is not None:
                try:
                    msg.close()
                except Exception:
                    pass
    
    @staticmethod
    def extract_eml(file_path: str) -> Optional[Dict]:
        """Extract email data from .eml file"""
        try:
            with open(file_path, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)
            attachments = []
            for part in msg.walk():
                if part.is_multipart():
                    continue
                disposition = part.get_content_disposition()
                filename = part.get_filename()
                if disposition not in ("attachment", "inline") and not filename:
                    continue
                payload = part.get_payload(decode=True)
                size_bytes = len(payload) if payload is not None else None
                attachments.append(
                    {
                        "filename": filename or "unnamed_attachment",
                        "content_type": part.get_content_type() or "",
                        "size_bytes": size_bytes,
                    }
                )
            body_part = msg.get_body(preferencelist=('plain', 'html'))
            body_text = body_part.get_content() if body_part else ''

            return {
                'subject': msg.get('subject', '(No Subject)'),
                'from': msg.get('from', ''),
                'to': msg.get('to', ''),
                'cc': msg.get('cc', ''),
                'date': msg.get('date'),
                'body': body_text,
                'attachments': attachments,
            }
        except Exception as e:
            print(f"Error extracting .eml file {file_path}: {e}")
            return None


class EmailThreader:
    """Groups emails into conversation threads"""
    
    @staticmethod
    def normalize_subject(subject: str) -> str:
        """Normalize email subject by removing RE:, FW:, etc."""
        if not subject:
            return ""
        
        # Remove prefixes like RE:, FW:, FWD:, etc.
        subject = re.sub(r'^(RE|FW|FWD):\s*', '', subject, flags=re.IGNORECASE)
        subject = re.sub(r'\s+', ' ', subject).strip()
        return subject.lower()

    @staticmethod
    def normalize_date(date_value) -> datetime:
        """Normalize date values to naive UTC datetimes for safe sorting."""
        if isinstance(date_value, datetime):
            parsed = date_value
        elif isinstance(date_value, str) and date_value.strip():
            if not HAS_DATEUTIL:
                raise RuntimeError(
                    "python-dateutil is required for email date parsing. "
                    "Install it with: pip install python-dateutil"
                )
            try:
                parsed = date_parser.parse(date_value)
            except (ValueError, TypeError, OverflowError):
                return datetime.min
        else:
            return datetime.min

        if parsed.tzinfo is not None:
            parsed = parsed.astimezone(timezone.utc).replace(tzinfo=None)
        return parsed
    
    def group_emails(self, email_data: List[Dict]) -> Dict[str, List[Dict]]:
        """
        Group emails into threads by normalized subject
        
        Args:
            email_data: List of dicts with 'subject', 'date', 'file_path', etc.
            
        Returns:
            Dict mapping thread_id to list of email dicts
        """
        threads = defaultdict(list)
        
        for email in email_data:
            normalized_subject = self.normalize_subject(email.get('subject', ''))
            thread_key = normalized_subject or f"no_subject_{email.get('file_path', '')}"
            threads[thread_key].append(email)
        
        # Sort emails within each thread by date
        for thread_emails in threads.values():
            thread_emails.sort(key=lambda e: self.normalize_date(e.get('date')))
        
        return dict(threads)


class FolderAnalyzer:
    """Analyzes folder structure and determines grouping strategy"""
    
    @staticmethod
    def analyze_structure(
        root_path: str,
        exclude_paths: Optional[List[str]] = None,
    ) -> Dict[str, List[str]]:
        """
        Analyze folder structure and group files by parent folder
        
        Args:
            root_path: Root directory to analyze
            exclude_paths: Optional list of directories to exclude from traversal
            
        Returns:
            Dict mapping group_name to list of file paths
        """
        groups = defaultdict(list)
        excluded = [
            os.path.normcase(os.path.abspath(path))
            for path in (exclude_paths or [])
            if path
        ]

        def is_excluded(candidate_path: str) -> bool:
            candidate = os.path.normcase(os.path.abspath(candidate_path))
            for excluded_path in excluded:
                if candidate == excluded_path or candidate.startswith(excluded_path + os.sep):
                    return True
            return False

        for dirpath, dirnames, filenames in os.walk(root_path):
            if is_excluded(dirpath):
                dirnames[:] = []
                continue

            dirnames[:] = [
                dirname
                for dirname in dirnames
                if not is_excluded(os.path.join(dirpath, dirname))
            ]

            for filename in filenames:
                file_path = os.path.join(dirpath, filename)
                
                # Determine group name from folder structure
                rel_path = os.path.relpath(dirpath, root_path)
                
                if rel_path == '.':
                    # Files in root directory
                    group_name = 'root'
                else:
                    # Use first subfolder as group name
                    group_name = rel_path.split(os.sep)[0]
                
                groups[group_name].append(file_path)
        
        return dict(groups)


class ZipArchiveProcessor:
    """Extract ZIP archives safely with truncation for long entry names."""

    @staticmethod
    def _safe_member_path(member_name: str) -> Optional[str]:
        if not member_name:
            return None

        normalized = member_name.replace("\\", "/")
        if normalized.startswith("/"):
            return None
        if re.match(r"^[A-Za-z]:", normalized):
            return None

        had_trailing_slash = normalized.endswith("/")
        parts = []
        for part in normalized.split("/"):
            if part in ("", "."):
                continue
            if part == "..":
                return None
            parts.append(part)

        if not parts:
            return None

        safe_path = "/".join(parts)
        if had_trailing_slash:
            safe_path += "/"
        return safe_path

    @staticmethod
    def _truncate_leaf_name(name: str, max_len: int, include_ext: bool) -> str:
        if max_len <= 0:
            return name

        base, ext = os.path.splitext(name)
        if include_ext:
            allowed_base = max(1, max_len - len(ext))
        else:
            allowed_base = max_len

        if len(base) > allowed_base:
            base = base[:allowed_base]
        return base + ext

    def _unique_path(
        self,
        candidate_path: str,
        used_paths: Set[str],
        max_len: int,
        include_ext: bool,
    ) -> Tuple[str, bool]:
        if candidate_path not in used_paths:
            return candidate_path, False

        parts = candidate_path.split("/")
        leaf_name = parts[-1]
        dir_prefix = "/".join(parts[:-1])
        base, ext = os.path.splitext(leaf_name)

        for counter in range(1, 100000):
            suffix = f"_{counter}"
            if max_len > 0:
                if include_ext:
                    allowed_base = max(1, max_len - len(ext) - len(suffix))
                else:
                    allowed_base = max(1, max_len - len(suffix))
                base_for_counter = base[:allowed_base]
            else:
                base_for_counter = base

            next_leaf = f"{base_for_counter}{suffix}{ext}"
            next_candidate = (
                f"{dir_prefix}/{next_leaf}" if dir_prefix else next_leaf
            )
            if next_candidate not in used_paths:
                return next_candidate, True

        raise RuntimeError("Unable to resolve ZIP filename collisions.")

    def extract_archive(
        self,
        zip_path: str,
        target_root: str,
        max_len: int,
        include_ext: bool,
        depth: int,
        depth_limit: int,
        warnings: Optional[List[Dict]],
        max_extract_bytes: int = 0,
    ) -> Dict:
        stats = {
            'archives_extracted': 0,
            'archives_failed': 0,
            'entries_total': 0,
            'entries_extracted': 0,
            'entries_renamed': 0,
            'entries_skipped_unsafe_path': 0,
            'nested_archives_extracted': 0,
            'nested_archives_skipped_depth': 0,
            'extracted_files': [],
        }

        used_paths: Set[str] = set()
        nested_archives: List[str] = []
        total_extracted_bytes = 0
        budget_exceeded = False

        try:
            with zipfile.ZipFile(zip_path, 'r') as zip_handle:
                stats['archives_extracted'] += 1

                for entry in zip_handle.infolist():
                    # ZIP bomb protection: skip if budget exceeded
                    if budget_exceeded:
                        break
                    if max_extract_bytes > 0 and entry.file_size > 0:
                        ratio = entry.file_size / max(entry.compress_size, 1)
                        if ratio > 100:
                            _record_warning(
                                warnings,
                                'zip_entry_suspicious_ratio',
                                'ZIP entry has suspicious compression ratio (possible zip bomb); skipping',
                                archive=zip_path,
                                entry=entry.filename,
                                ratio=round(ratio, 1),
                            )
                            continue
                        if total_extracted_bytes + entry.file_size > max_extract_bytes:
                            _record_warning(
                                warnings,
                                'zip_extraction_budget_exceeded',
                                'ZIP extraction stopped: total extracted size would exceed safety budget',
                                archive=zip_path,
                                budget_bytes=max_extract_bytes,
                                extracted_so_far=total_extracted_bytes,
                            )
                            budget_exceeded = True
                            break

                    safe_member = self._safe_member_path(entry.filename)
                    if safe_member is None:
                        if not entry.is_dir():
                            stats['entries_total'] += 1
                            stats['entries_skipped_unsafe_path'] += 1
                            _record_warning(
                                warnings,
                                'zip_entry_skipped_unsafe_path',
                                'Skipped ZIP entry with unsafe path',
                                archive=zip_path,
                                entry=entry.filename,
                            )
                        continue

                    if safe_member.endswith("/"):
                        continue

                    stats['entries_total'] += 1
                    member_parts = safe_member.split("/")
                    leaf_name = member_parts[-1]
                    parent_parts = member_parts[:-1]

                    truncated_leaf = self._truncate_leaf_name(
                        leaf_name,
                        max_len,
                        include_ext,
                    )
                    renamed = truncated_leaf != leaf_name
                    candidate_rel = "/".join(parent_parts + [truncated_leaf]) if parent_parts else truncated_leaf
                    unique_rel, renamed_by_collision = self._unique_path(
                        candidate_rel,
                        used_paths,
                        max_len,
                        include_ext,
                    )
                    if renamed or renamed_by_collision:
                        stats['entries_renamed'] += 1
                    used_paths.add(unique_rel)

                    target_path = os.path.join(target_root, *unique_rel.split("/"))
                    os.makedirs(os.path.dirname(target_path), exist_ok=True)

                    try:
                        with zip_handle.open(entry) as source_handle:
                            with open(target_path, 'wb') as target_handle:
                                shutil.copyfileobj(source_handle, target_handle)
                        stats['entries_extracted'] += 1
                        total_extracted_bytes += entry.file_size
                    except Exception as exc:
                        _record_warning(
                            warnings,
                            'zip_extract_failed',
                            'Failed to extract ZIP entry; skipping entry',
                            archive=zip_path,
                            entry=entry.filename,
                            error=str(exc),
                        )
                        continue

                    if target_path.lower().endswith('.zip'):
                        nested_archives.append(target_path)
                    else:
                        stats['extracted_files'].append(target_path)
        except Exception as exc:
            stats['archives_failed'] += 1
            _record_warning(
                warnings,
                'zip_extract_failed',
                'Failed to extract ZIP archive; skipping archive',
                archive=zip_path,
                error=str(exc),
            )
            return stats

        if nested_archives:
            if depth < depth_limit:
                for nested_archive in nested_archives:
                    nested_target = os.path.join(target_root, f"_nested_{uuid.uuid4().hex}")
                    os.makedirs(nested_target, exist_ok=True)
                    nested_stats = self.extract_archive(
                        nested_archive,
                        nested_target,
                        max_len=max_len,
                        include_ext=include_ext,
                        depth=depth + 1,
                        depth_limit=depth_limit,
                        warnings=warnings,
                        max_extract_bytes=max_extract_bytes,
                    )
                    stats['archives_extracted'] += nested_stats['archives_extracted']
                    stats['archives_failed'] += nested_stats['archives_failed']
                    stats['entries_total'] += nested_stats['entries_total']
                    stats['entries_extracted'] += nested_stats['entries_extracted']
                    stats['entries_renamed'] += nested_stats['entries_renamed']
                    stats['entries_skipped_unsafe_path'] += nested_stats['entries_skipped_unsafe_path']
                    stats['nested_archives_extracted'] += 1  # count this immediate nested archive
                    stats['nested_archives_extracted'] += nested_stats['nested_archives_extracted']
                    stats['nested_archives_skipped_depth'] += nested_stats['nested_archives_skipped_depth']
                    stats['extracted_files'].extend(nested_stats['extracted_files'])
            else:
                for nested_archive in nested_archives:
                    stats['nested_archives_skipped_depth'] += 1
                    _record_warning(
                        warnings,
                        'zip_nested_depth_exceeded',
                        'Nested ZIP archive skipped due to depth limit',
                        archive=nested_archive,
                        depth=depth,
                        depth_limit=depth_limit,
                    )

        return stats


class MergeOrchestrator:
    """Coordinates the entire merging process"""
    
    def __init__(
        self,
        max_file_size_kb=102400,
        max_output_files=300,
        process_pdfs=True,
        process_docx=True,
        process_emails=True,
        process_zip_archives=True,
        zip_max_filename_length=50,
        zip_include_extension_in_limit=True,
        zip_nested_depth_limit=1,
        output_layout_mode="structured",
        processed_subdir="processed",
        unprocessed_subdir="unprocessed",
        failed_subdir="failed",
        logs_subdir="logs",
        word_progress_interval=10,
        enable_detailed_logging=True,
        log_privacy_mode="redacted",
        relocate_unsupported_input_files=True,
        unsupported_input_action="copy",
        failed_file_artifact_action="copy",
        unprocessed_include_source_files=True,
        failed_include_artifacts=True,
        email_output_mode="size_batched",
        email_max_output_file_mb=25,
        email_include_attachment_index=True,
        email_batch_name_prefix="emails_batch",
        zip_max_extract_bytes=2 * 1024 ** 3,  # 2 GB default extraction budget
        word_convert_timeout_seconds=120,
    ):
        self.max_file_size_kb = max_file_size_kb
        self.pdf_merger = PDFMerger(max_file_size_kb)
        self.email_extractor = EmailExtractor()
        self.email_threader = EmailThreader()
        self.folder_analyzer = FolderAnalyzer()
        self.zip_archive_processor = ZipArchiveProcessor()
        self.word_converter_factory = WordToPdfConverter
        self.max_output_files = max_output_files
        self.process_pdfs = process_pdfs
        self.process_docx = process_docx
        self.process_emails = process_emails
        self.process_zip_archives = process_zip_archives
        self.zip_max_filename_length = zip_max_filename_length
        self.zip_include_extension_in_limit = zip_include_extension_in_limit
        self.zip_nested_depth_limit = zip_nested_depth_limit
        self.output_layout_mode = output_layout_mode
        self.processed_subdir = processed_subdir
        self.unprocessed_subdir = unprocessed_subdir
        self.failed_subdir = failed_subdir
        self.logs_subdir = logs_subdir
        self.word_progress_interval = max(1, int(word_progress_interval))
        self.enable_detailed_logging = enable_detailed_logging
        self.log_privacy_mode = log_privacy_mode
        self.relocate_unsupported_input_files = relocate_unsupported_input_files
        self.unsupported_input_action = unsupported_input_action
        self.failed_file_artifact_action = failed_file_artifact_action
        self.unprocessed_include_source_files = unprocessed_include_source_files
        self.failed_include_artifacts = failed_include_artifacts
        self.email_output_mode = email_output_mode
        self.email_max_output_file_mb = max(1, int(email_max_output_file_mb))
        self.email_include_attachment_index = email_include_attachment_index
        self.email_batch_name_prefix = email_batch_name_prefix
        self.zip_max_extract_bytes = int(zip_max_extract_bytes)
        self.word_convert_timeout_seconds = max(10, int(word_convert_timeout_seconds))

    def merge_documents(
        self,
        input_path: str,
        output_path: str,
        progress_callback=None,
        event_callback: Optional[Callable[[Dict[str, Any]], None]] = None,
        cancel_event: Optional[threading.Event] = None,
    ) -> Dict:
        """
        Main entry point for document merging

        Args:
            input_path: Input directory containing documents
            output_path: Output directory for merged files
            progress_callback: Optional callback function(current, total, message)
            cancel_event: Optional threading.Event; set to request graceful cancellation

        Returns:
            Dict with merge statistics
        """
        print(f"\nAnalyzing folder structure: {input_path}")

        os.makedirs(output_path, exist_ok=True)
        if self.output_layout_mode == "structured":
            processed_dir = os.path.join(output_path, self.processed_subdir)
            unprocessed_dir = os.path.join(output_path, self.unprocessed_subdir)
            failed_dir = os.path.join(output_path, self.failed_subdir)
            logs_dir = os.path.join(output_path, self.logs_subdir)
        else:
            processed_dir = output_path
            unprocessed_dir = os.path.join(output_path, self.unprocessed_subdir)
            failed_dir = os.path.join(output_path, self.failed_subdir)
            logs_dir = os.path.join(output_path, self.logs_subdir)

        for path in (processed_dir, unprocessed_dir, failed_dir, logs_dir):
            os.makedirs(path, exist_ok=True)

        run_id = datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:8]
        run_logger = RunLogger(
            logs_dir=logs_dir,
            run_id=run_id,
            enabled=self.enable_detailed_logging,
            privacy_mode=self.log_privacy_mode,
            event_callback=event_callback,
        )

        staged_input_root = None
        working_input_path = input_path
        output_files: List[str] = []
        warnings: List[Dict] = []
        errors: List[Dict] = []
        output_to_sources: Dict[str, List[str]] = {}
        moved_unprocessed: List[Dict] = []
        unprocessed_files: List[Dict] = []
        failed_files: List[Dict] = []
        skipped_files: List[Dict] = []
        fatal_exception = None
        fatal_error_message = ""
        warning_cursor = 0
        failed_artifacts_total = 0

        word_conversion_summary = {
            'attempted': 0,
            'converted': 0,
            'failed': 0,
        }
        zip_processing_summary = {
            'archives_found': 0,
            'archives_extracted': 0,
            'archives_failed': 0,
            'entries_total': 0,
            'entries_extracted': 0,
            'entries_renamed': 0,
            'entries_skipped_unsafe_path': 0,
            'nested_archives_extracted': 0,
            'nested_archives_skipped_depth': 0,
        }

        total_input_files = 0
        file_count = 0
        zip_temp_dirs: List[str] = []
        manifest: Dict = {}
        email_summary = {
            'parsed_total': 0,
            'failed_total': 0,
            'threads_total': 0,
            'batches_total': 0,
            'output_total_bytes': 0,
            'attachment_refs_total': 0,
            'batch_to_threads': {},
        }

        try:
            if os.path.isfile(input_path):
                if not input_path.lower().endswith('.zip'):
                    raise RuntimeError("Input path must be a folder or .zip file.")
                staged_input_root = _make_writable_temp_dir(prefix="single_zip_input_")
                with _active_temp_dirs_lock:
                    _active_temp_dirs.add(staged_input_root)
                staged_archive = os.path.join(staged_input_root, os.path.basename(input_path))
                shutil.copy2(input_path, staged_archive)
                working_input_path = staged_input_root
            elif not os.path.isdir(input_path):
                raise RuntimeError("Input path must exist and be accessible.")

            input_abs = os.path.normcase(os.path.abspath(working_input_path))
            output_abs = os.path.normcase(os.path.abspath(output_path))
            exclude_paths = []
            if output_abs == input_abs or output_abs.startswith(input_abs + os.sep):
                exclude_paths.append(output_path)
                print(f"Excluding output folder from scan: {output_path}")

            groups = self.folder_analyzer.analyze_structure(working_input_path, exclude_paths=exclude_paths)
            print(f"Found {len(groups)} groups to process")
            run_logger.log("info", "groups_analyzed", "Folder analysis complete", group_count=len(groups))
            warning_cursor = self._sync_warning_events(warnings, warning_cursor, run_logger)

            total_input_files = sum(len(files) for files in groups.values())
            if staged_input_root:
                zip_temp_dirs.append(staged_input_root)

            groups, group_file_weights, extracted_zip_temp_dirs, zip_group_meta = self._prepare_groups_with_zip_expansion(
                groups,
                warnings=warnings,
                zip_processing_summary=zip_processing_summary,
            )
            zip_temp_dirs.extend(extracted_zip_temp_dirs)
            warning_cursor = self._sync_warning_events(warnings, warning_cursor, run_logger)

            for group_name, files in sorted(groups.items()):
                if cancel_event is not None and cancel_event.is_set():
                    run_logger.log("warning", "run_cancelled", "Merge cancelled by user")
                    break
                if not files:
                    continue

                print(f"\nProcessing group: {group_name}")
                run_logger.log("info", "group_start", "Processing group", group=group_name, file_count=len(files))
                group_files = sorted(files)

                supported_files = [f for f in group_files if self._is_supported_processable_file(f)]
                unsupported_files = [f for f in group_files if f not in supported_files]
                if unsupported_files and self.unprocessed_include_source_files:
                    if group_name in zip_group_meta:
                        relocated = self._relocate_unsupported_files(
                            files_to_relocate=unsupported_files,
                            target_root=unprocessed_dir,
                            target_prefix="",
                            base_path=zip_group_meta[group_name]["extraction_root"],
                            action="move",
                            reason="unsupported_zip_file_moved",
                            origin="zip_extract",
                            stage="classification",
                            warnings=warnings,
                            run_logger=run_logger,
                            success_event="unsupported_zip_file_relocated",
                            flatten=True,
                        )
                        unprocessed_files.extend(relocated)
                        moved_unprocessed.extend(
                            {
                                "source": item["source"],
                                "destination": item["destination"],
                                "reason": item["reason"],
                            }
                            for item in relocated
                        )
                    elif self.relocate_unsupported_input_files:
                        relocated = self._relocate_unsupported_files(
                            files_to_relocate=unsupported_files,
                            target_root=unprocessed_dir,
                            target_prefix="",
                            base_path=working_input_path,
                            action=self.unsupported_input_action,
                            reason="unsupported_input_file_relocated",
                            origin="input",
                            stage="classification",
                            warnings=warnings,
                            run_logger=run_logger,
                            success_event="unsupported_input_file_relocated",
                            flatten=True,
                        )
                        unprocessed_files.extend(relocated)
                group_files = supported_files

                pdfs = [f for f in group_files if f.lower().endswith('.pdf')]
                word_docs = [f for f in group_files if f.lower().endswith(('.docx', '.doc'))]
                emails = [f for f in group_files if f.lower().endswith(('.msg', '.eml'))]

                if pdfs and self.process_pdfs:
                    print(f"  Merging {len(pdfs)} PDF files...")
                    run_logger.log("info", "pdf_merge_start", "Starting PDF merge", group=group_name, count=len(pdfs))
                    required_pdf_outputs = self.pdf_merger.estimate_batch_count(pdfs)
                    self._ensure_output_capacity(
                        required_pdf_outputs,
                        len(output_files),
                        f"group '{group_name}' PDF files",
                    )
                    pdf_outputs = self.pdf_merger.merge_pdfs(
                        pdfs,
                        processed_dir,
                        group_name,
                        warnings=warnings,
                    )
                    output_files.extend(pdf_outputs)
                    run_logger.log("info", "pdf_merge_end", "Completed PDF merge", group=group_name, outputs=len(pdf_outputs))
                    warning_cursor = self._sync_warning_events(warnings, warning_cursor, run_logger)

                if word_docs and self.process_docx:
                    print(f"  Processing {len(word_docs)} Word document files...")
                    run_logger.log("info", "word_convert_start", "Starting Word conversion", group=group_name, count=len(word_docs))

                    def word_progress_update(message: str) -> None:
                        _safe_progress(progress_callback, file_count, total_input_files, message)

                    doc_outputs, doc_output_to_sources, conversion_summary = self._process_word_documents(
                        word_docs,
                        processed_dir,
                        group_name,
                        len(output_files),
                        warnings=warnings,
                        progress_callback=word_progress_update,
                        progress_interval=self.word_progress_interval,
                        run_logger=run_logger,
                    )
                    output_files.extend(doc_outputs)
                    output_to_sources.update(doc_output_to_sources)
                    word_conversion_summary['attempted'] += conversion_summary['attempted']
                    word_conversion_summary['converted'] += conversion_summary['converted']
                    word_conversion_summary['failed'] += conversion_summary['failed']
                    run_logger.log("info", "word_convert_end", "Completed Word conversion", group=group_name, outputs=len(doc_outputs))
                    warning_cursor = self._sync_warning_events(warnings, warning_cursor, run_logger)

                if emails and self.process_emails:
                    print(f"  Processing {len(emails)} email files...")
                    run_logger.log("info", "email_thread_start", "Starting email threading", group=group_name, count=len(emails))
                    email_threads, email_extract_stats = self._prepare_email_threads(
                        emails,
                        warnings=warnings,
                    )
                    email_summary['parsed_total'] += email_extract_stats.get('parsed_total', 0)
                    email_summary['failed_total'] += email_extract_stats.get('failed_total', 0)
                    email_summary['attachment_refs_total'] += email_extract_stats.get('attachment_refs_total', 0)
                    warning_cursor = self._sync_warning_events(warnings, warning_cursor, run_logger)
                    email_outputs, email_batch_stats = self._write_email_outputs(
                        threads=email_threads,
                        output_path=processed_dir,
                        group_name=group_name,
                        current_output_count=len(output_files),
                        warnings=warnings,
                        run_logger=run_logger,
                    )
                    output_files.extend(email_outputs)
                    email_summary['threads_total'] += email_batch_stats.get('threads_total', 0)
                    email_summary['batches_total'] += email_batch_stats.get('batches_total', 0)
                    email_summary['output_total_bytes'] += email_batch_stats.get('output_total_bytes', 0)
                    email_summary['batch_to_threads'].update(email_batch_stats.get('batch_to_threads', {}))
                    run_logger.log("info", "email_thread_end", "Completed email threading", group=group_name, outputs=len(email_outputs))
                    warning_cursor = self._sync_warning_events(warnings, warning_cursor, run_logger)

                file_count += group_file_weights.get(group_name, 0)
                _safe_progress(progress_callback, file_count, total_input_files, f"Processed {group_name}")
                run_logger.log("info", "group_end", "Finished group", group=group_name, processed=file_count, total=total_input_files)
                warning_cursor = self._sync_warning_events(warnings, warning_cursor, run_logger)

            failed_files, skipped_files = self._collect_file_outcomes_from_warnings(warnings)
            failed_artifacts_total = self._materialize_failed_artifacts(
                failed_files=failed_files,
                failed_root=failed_dir,
                input_root=working_input_path,
                action=self.failed_file_artifact_action,
                include_artifacts=self.failed_include_artifacts,
                warnings=warnings,
                run_logger=run_logger,
            )
            warning_cursor = self._sync_warning_events(warnings, warning_cursor, run_logger)
        except Exception as exc:
            fatal_exception = exc
            fatal_error_message = str(exc)
            errors.append(
                {
                    "code": "infra_unhandled_error",
                    "message": "Fatal processing error stopped the run",
                    "error": str(exc),
                    "traceback": traceback.format_exc(),
                }
            )
            run_logger.log("error", "fatal_error", "Fatal processing error", error=str(exc))
        finally:
            if not failed_files and not skipped_files:
                collected_failed, collected_skipped = self._collect_file_outcomes_from_warnings(warnings)
                if not failed_files:
                    failed_files = collected_failed
                if not skipped_files:
                    skipped_files = collected_skipped
            if failed_files and failed_artifacts_total == 0:
                failed_artifacts_total = self._materialize_failed_artifacts(
                    failed_files=failed_files,
                    failed_root=failed_dir,
                    input_root=working_input_path,
                    action=self.failed_file_artifact_action,
                    include_artifacts=self.failed_include_artifacts,
                    warnings=warnings,
                    run_logger=run_logger,
                )
                warning_cursor = self._sync_warning_events(warnings, warning_cursor, run_logger)

            manifest_path = os.path.join(processed_dir, 'merge_manifest.json')
            manifest = {
                'timestamp': datetime.now().isoformat(),
                'run_id': run_id,
                'input_path': input_path,
                'total_input_files': total_input_files,
                'total_output_files': len(output_files),
                'output_files': output_files,
                'limits': {
                    'max_file_size_kb': self.max_file_size_kb,
                    'max_output_files': self.max_output_files,
                },
                'paths': {
                    'processed_dir': processed_dir,
                    'unprocessed_dir': unprocessed_dir,
                    'failed_dir': failed_dir,
                    'logs_dir': logs_dir,
                },
                'logs': {
                    'text_log': run_logger.text_log_path,
                    'jsonl_log': run_logger.jsonl_log_path,
                },
                'summary': {
                    'input_files_total': total_input_files,
                    'processed_outputs_total': len(output_files),
                    'moved_unprocessed_total': len(moved_unprocessed),
                    'unprocessed_relocated_total': len(unprocessed_files),
                    'failed_files_total': len(failed_files),
                    'failed_artifacts_total': failed_artifacts_total,
                    'skipped_files_total': len(skipped_files),
                    'warnings_total': len(warnings),
                    'errors_total': len(errors),
                },
                'files': {
                    'processed_outputs': output_files,
                    'moved_unprocessed': moved_unprocessed,
                    'unprocessed': unprocessed_files,
                    'failed': failed_files,
                    'skipped': skipped_files,
                },
            }

            if warnings:
                manifest['warnings'] = warnings
            if errors:
                manifest['errors'] = errors
            if output_to_sources:
                manifest['output_to_sources'] = output_to_sources
            if word_conversion_summary['attempted'] > 0:
                manifest['word_conversion'] = word_conversion_summary
            if zip_processing_summary['archives_found'] > 0:
                manifest['zip_processing'] = zip_processing_summary
            if email_summary['parsed_total'] > 0 or email_summary['failed_total'] > 0:
                manifest['emails'] = email_summary

            try:
                with open(manifest_path, 'w', encoding='utf-8') as f:
                    json.dump(manifest, f, indent=2)
                run_logger.log("info", "manifest_written", "Wrote merge manifest", path=manifest_path)
            except Exception as manifest_exc:
                run_logger.log(
                    "warning",
                    "manifest_write_failed",
                    f"Could not write merge_manifest.json: {manifest_exc}",
                    path=manifest_path,
                )
                manifest['manifest_write_error'] = str(manifest_exc)
            finally:
                run_logger.close()
                for temp_dir in zip_temp_dirs:
                    shutil.rmtree(temp_dir, ignore_errors=True)
                # Deregister from atexit safety net after normal cleanup.
                with _active_temp_dirs_lock:
                    for temp_dir in zip_temp_dirs:
                        _active_temp_dirs.discard(temp_dir)

        if fatal_exception is not None:
            raise RuntimeError(
                f"{fatal_error_message}\n"
                f"Run log: {run_logger.text_log_path}\n"
                f"Manifest: {os.path.join(processed_dir, 'merge_manifest.json')}"
            ) from fatal_exception

        return manifest

    @staticmethod
    def _sanitize_group_component(value: str) -> str:
        normalized = re.sub(r'[^A-Za-z0-9_-]+', '_', value).strip('_')
        return normalized or 'zip'

    @classmethod
    def _allocate_zip_group_name(
        cls,
        group_name: str,
        zip_path: str,
        used_group_names: Set[str],
    ) -> str:
        stem = os.path.splitext(os.path.basename(zip_path))[0]
        zip_component = cls._sanitize_group_component(stem)
        base_name = f"{group_name}_{zip_component}" if group_name else zip_component

        if base_name not in used_group_names:
            return base_name

        for counter in range(2, 100000):
            candidate = f"{base_name}_{counter}"
            if candidate not in used_group_names:
                return candidate

        raise RuntimeError("Unable to generate a unique group name for ZIP archive.")

    @staticmethod
    def _merge_zip_stats(accumulator: Dict[str, int], update: Dict[str, int]) -> None:
        for key in (
            'archives_extracted',
            'archives_failed',
            'entries_total',
            'entries_extracted',
            'entries_renamed',
            'entries_skipped_unsafe_path',
            'nested_archives_extracted',
            'nested_archives_skipped_depth',
        ):
            accumulator[key] += update.get(key, 0)

    def _prepare_groups_with_zip_expansion(
        self,
        groups: Dict[str, List[str]],
        warnings: Optional[List[Dict]],
        zip_processing_summary: Dict[str, int],
    ) -> Tuple[Dict[str, List[str]], Dict[str, int], List[str], Dict[str, Dict[str, str]]]:
        expanded_groups: Dict[str, List[str]] = defaultdict(list)
        group_file_weights: Dict[str, int] = defaultdict(int)
        temp_dirs: List[str] = []
        zip_group_meta: Dict[str, Dict[str, str]] = {}
        used_group_names: Set[str] = set(groups.keys())

        for group_name, files in sorted(groups.items()):
            sorted_files = sorted(files)
            zip_files = [path for path in sorted_files if path.lower().endswith('.zip')]
            non_zip_files = [path for path in sorted_files if not path.lower().endswith('.zip')]

            if non_zip_files:
                expanded_groups[group_name].extend(non_zip_files)
                group_file_weights[group_name] += len(non_zip_files)

            if not self.process_zip_archives:
                group_file_weights[group_name] += len(zip_files)
                continue

            for zip_file in zip_files:
                zip_processing_summary['archives_found'] += 1
                zip_group_name = self._allocate_zip_group_name(
                    group_name,
                    zip_file,
                    used_group_names,
                )
                used_group_names.add(zip_group_name)
                group_file_weights[zip_group_name] += 1

                try:
                    extraction_root = _make_writable_temp_dir(prefix=f"zip_extract_{zip_group_name}_")
                except Exception as exc:
                    zip_processing_summary['archives_failed'] += 1
                    _record_warning(
                        warnings,
                        'zip_extract_failed',
                        'Failed to prepare extraction directory; skipping archive',
                        archive=zip_file,
                        error=str(exc),
                    )
                    continue

                temp_dirs.append(extraction_root)
                with _active_temp_dirs_lock:
                    _active_temp_dirs.add(extraction_root)
                extract_result = self.zip_archive_processor.extract_archive(
                    zip_file,
                    extraction_root,
                    max_len=self.zip_max_filename_length,
                    include_ext=self.zip_include_extension_in_limit,
                    depth=0,
                    depth_limit=self.zip_nested_depth_limit,
                    warnings=warnings,
                    max_extract_bytes=self.zip_max_extract_bytes,
                )
                self._merge_zip_stats(zip_processing_summary, extract_result)

                extracted_files = extract_result.get('extracted_files', [])
                if extracted_files:
                    expanded_groups[zip_group_name].extend(extracted_files)
                    zip_group_meta[zip_group_name] = {
                        "source_archive": zip_file,
                        "extraction_root": extraction_root,
                    }
                else:
                    _record_warning(
                        warnings,
                        'zip_empty_after_extraction',
                        'ZIP archive did not contain extractable files',
                        archive=zip_file,
                    )

        return dict(expanded_groups), dict(group_file_weights), temp_dirs, zip_group_meta

    @staticmethod
    def _is_supported_processable_file(file_path: str) -> bool:
        supported = ('.pdf', '.docx', '.doc', '.msg', '.eml')
        return file_path.lower().endswith(supported)

    @staticmethod
    def _relative_path_under(source_path: str, base_path: str) -> str:
        try:
            relative = os.path.relpath(source_path, base_path)
            if relative.startswith(".."):
                return os.path.basename(source_path)
            return relative
        except ValueError:
            return os.path.basename(source_path)

    @staticmethod
    def _to_windows_long_path(path: str) -> str:
        if os.name != "nt":
            return path
        normalized = os.path.abspath(path).replace("/", "\\")
        if normalized.startswith("\\\\?\\"):
            return normalized
        return "\\\\?\\" + normalized

    @classmethod
    def _path_is_file(cls, path: str) -> bool:
        if os.path.isfile(path):
            return True
        if os.name == "nt":
            return os.path.isfile(cls._to_windows_long_path(path))
        return False

    @staticmethod
    def _truncate_leaf_name(name: str, max_len: int = 120) -> str:
        if max_len <= 0 or len(name) <= max_len:
            return name
        base, ext = os.path.splitext(name)
        keep_base = max(1, max_len - len(ext))
        return base[:keep_base] + ext

    @classmethod
    def _copy_or_move_file(cls, source: str, destination: str, action: str) -> None:
        src = source
        dst = destination
        if os.name == "nt":
            src = cls._to_windows_long_path(source)
            dst = cls._to_windows_long_path(destination)
        if action == "move":
            shutil.move(src, dst)
        else:
            shutil.copy2(src, dst)

    @staticmethod
    def _ensure_unique_destination(path: str) -> str:
        if not os.path.exists(path):
            return path
        directory, filename = os.path.split(path)
        base, ext = os.path.splitext(filename)
        for index in range(1, 100000):
            candidate = os.path.join(directory, f"{base}_{index}{ext}")
            if not os.path.exists(candidate):
                return candidate
        raise RuntimeError("Unable to allocate destination filename.")

    @staticmethod
    def _sync_warning_events(
        warnings: List[Dict],
        cursor: int,
        run_logger: Optional[RunLogger],
    ) -> int:
        if not run_logger:
            return len(warnings)
        while cursor < len(warnings):
            warning = warnings[cursor]
            code = warning.get("code", "warning")
            message = warning.get("message", "")
            context = {
                key: value
                for key, value in warning.items()
                if key not in {"code", "message"}
            }
            run_logger.log("warning", code, message, **context)
            cursor += 1
        return cursor

    def _relocate_unsupported_files(
        self,
        files_to_relocate: List[str],
        target_root: str,
        target_prefix: str,
        base_path: str,
        action: str,
        reason: str,
        origin: str,
        stage: str,
        warnings: Optional[List[Dict]],
        run_logger: Optional[RunLogger] = None,
        success_event: str = "unsupported_file_relocated",
        flatten: bool = False,
    ) -> List[Dict]:
        relocated_entries: List[Dict] = []
        normalized_action = "move" if str(action).lower() == "move" else "copy"

        for source_path in sorted(files_to_relocate):
            if flatten:
                safe_name = self._truncate_leaf_name(os.path.basename(source_path))
                destination = os.path.join(target_root, safe_name)
            else:
                relative = self._relative_path_under(source_path, base_path)
                destination = os.path.join(target_root, target_prefix, relative)
            os.makedirs(os.path.dirname(destination), exist_ok=True)
            destination = self._ensure_unique_destination(destination)

            try:
                self._copy_or_move_file(source_path, destination, normalized_action)
                entry = {
                    "source": source_path,
                    "destination": destination,
                    "action": normalized_action,
                    "reason": reason,
                    "origin": origin,
                    "stage": stage,
                }
                relocated_entries.append(entry)
                if run_logger:
                    run_logger.log(
                        "info",
                        success_event,
                        "Relocated unsupported file",
                        source=source_path,
                        destination=destination,
                        action=normalized_action,
                        reason=reason,
                        origin=origin,
                    )
            except Exception as exc:
                _record_warning(
                    warnings,
                    'unsupported_relocate_failed',
                    'Failed to relocate unsupported file',
                    file=source_path,
                    destination=destination,
                    action=normalized_action,
                    reason=reason,
                    origin=origin,
                    error=str(exc),
                )
        return relocated_entries

    def _materialize_failed_artifacts(
        self,
        failed_files: List[Dict],
        failed_root: str,
        input_root: str,
        action: str,
        include_artifacts: bool,
        warnings: Optional[List[Dict]],
        run_logger: Optional[RunLogger] = None,
    ) -> int:
        created = 0
        normalized_action = str(action).lower()
        if normalized_action not in {"copy", "move", "metadata_only"}:
            normalized_action = "copy"

        for item in failed_files:
            item["artifact_action"] = normalized_action
            item["artifact_status"] = "not_created"
            if not include_artifacts or normalized_action == "metadata_only":
                continue

            source = item.get("source")
            if not source or "::" in source:
                item["artifact_status"] = "source_missing"
                continue
            source = os.path.normpath(source)
            if not self._path_is_file(source):
                item["artifact_status"] = "source_missing"
                continue

            stage = item.get("stage") or "unknown"
            leaf_name = self._truncate_leaf_name(os.path.basename(source))
            destination = os.path.join(failed_root, stage, leaf_name)
            os.makedirs(os.path.dirname(destination), exist_ok=True)
            destination = self._ensure_unique_destination(destination)

            try:
                self._copy_or_move_file(source, destination, normalized_action)
                item["artifact_destination"] = destination
                item["artifact_status"] = "created"
                created += 1
                if run_logger:
                    run_logger.log(
                        "info",
                        "failed_artifact_created",
                        "Created failed file artifact",
                        source=source,
                        destination=destination,
                        stage=stage,
                        action=normalized_action,
                    )
            except Exception as exc:
                item["artifact_destination"] = destination
                item["artifact_status"] = "copy_failed"
                _record_warning(
                    warnings,
                    "failed_artifact_create_failed",
                    "Failed to create artifact for failed file",
                    file=source,
                    destination=destination,
                    stage=stage,
                    action=normalized_action,
                    error=str(exc),
                )
        return created

    @staticmethod
    def _collect_file_outcomes_from_warnings(
        warnings: List[Dict],
    ) -> Tuple[List[Dict], List[Dict]]:
        skip_codes = {
            'zip_entry_skipped_unsafe_path',
            'zip_nested_depth_exceeded',
            'zip_empty_after_extraction',
        }
        failed = []
        skipped = []
        seen = set()

        for warning in warnings:
            code = warning.get("code", "unknown_warning")
            message = warning.get("message", "")
            stage = code.split("_", 1)[0]

            source = warning.get("file")
            if source is None and warning.get("archive") and warning.get("entry"):
                source = f"{warning['archive']}::{warning['entry']}"
            if source is None:
                source = warning.get("archive") or warning.get("destination")
            if not source:
                continue

            item = {
                "source": source,
                "code": code,
                "message": message,
                "stage": stage,
            }
            key = (item["source"], item["code"], item["message"], item["stage"])
            if key in seen:
                continue
            seen.add(key)

            if code in skip_codes:
                skipped.append(item)
            else:
                failed.append(item)

        return failed, skipped

    def _process_word_documents(
        self,
        word_files: List[str],
        output_path: str,
        group_name: str,
        current_output_count: int,
        warnings: Optional[List[Dict]] = None,
        progress_callback=None,
        progress_interval: int = 10,
        run_logger: Optional[RunLogger] = None,
    ) -> Tuple[List[str], Dict[str, List[str]], Dict[str, int]]:
        """Convert .doc/.docx files to PDF, then merge converted PDFs."""
        word_files = sorted(word_files)
        conversion_summary = {
            'attempted': len(word_files),
            'converted': 0,
            'failed': 0,
        }

        available, reason = self.word_converter_factory.is_available()
        if not available:
            raise RuntimeError(
                "Word document processing requires Microsoft Word automation. "
                f"Details: {reason}"
            )

        conversion_dir = _make_writable_temp_dir(prefix=f"word_pdf_{group_name}_")
        converted_pdf_files: List[str] = []
        source_file_map: Dict[str, str] = {}
        bookmark_titles: Dict[str, str] = {}

        try:
            with self.word_converter_factory(warnings=warnings, timeout_seconds=self.word_convert_timeout_seconds) as converter:
                total_word_files = len(word_files)
                for index, source_file in enumerate(word_files, 1):
                    converted_pdf = os.path.join(conversion_dir, f"{uuid.uuid4().hex}.pdf")
                    converted = converter.convert_file(source_file, converted_pdf)
                    if converted:
                        conversion_summary['converted'] += 1
                        converted_pdf_files.append(converted_pdf)
                        source_file_map[converted_pdf] = source_file
                        bookmark_titles[converted_pdf] = os.path.basename(source_file)
                    else:
                        conversion_summary['failed'] += 1

                    if (index % max(1, progress_interval) == 0) or index == total_word_files:
                        message = (
                            f"Word conversion progress for {group_name}: "
                            f"{index}/{total_word_files} "
                            f"(converted={conversion_summary['converted']}, failed={conversion_summary['failed']})"
                        )
                        print(f"    {message}")
                        if run_logger:
                            run_logger.log(
                                "info",
                                "word_conversion_progress",
                                message,
                                group=group_name,
                                converted=conversion_summary['converted'],
                                failed=conversion_summary['failed'],
                                processed=index,
                                total=total_word_files,
                            )
                        _safe_progress(progress_callback, message)

            if not converted_pdf_files:
                _record_warning(
                    warnings,
                    'word_conversion_no_outputs',
                    'No Word documents could be converted to PDF in this group',
                    group=group_name,
                    attempted=len(word_files),
                )
                return [], {}, conversion_summary

            required_outputs = self.pdf_merger.estimate_batch_count(converted_pdf_files)
            self._ensure_output_capacity(
                required_outputs,
                current_output_count,
                f"group '{group_name}' Word document files",
            )

            output_to_sources: Dict[str, List[str]] = {}
            doc_outputs = self.pdf_merger.merge_pdfs(
                converted_pdf_files,
                output_path,
                group_name,
                warnings=warnings,
                output_label="documents",
                bookmark_titles=bookmark_titles,
                source_file_map=source_file_map,
                output_to_sources=output_to_sources,
            )
            summary_message = (
                f"Word conversion summary for {group_name}: "
                f"attempted={conversion_summary['attempted']}, "
                f"converted={conversion_summary['converted']}, "
                f"failed={conversion_summary['failed']}"
            )
            print(f"    {summary_message}")
            if run_logger:
                run_logger.log("info", "word_conversion_summary", summary_message, group=group_name)
            _safe_progress(progress_callback, summary_message)
            return doc_outputs, output_to_sources, conversion_summary
        finally:
            shutil.rmtree(conversion_dir, ignore_errors=True)

    def _ensure_output_capacity(
        self,
        required_outputs: int,
        current_output_count: int,
        context: str,
    ) -> None:
        """Validate that writing the next outputs will stay under the global limit."""
        if required_outputs <= 0:
            return

        remaining = self.max_output_files - current_output_count
        if required_outputs > remaining:
            raise RuntimeError(
                f"Output file limit exceeded before processing {context}: "
                f"requires {required_outputs} file(s), but only {remaining} slot(s) remain "
                f"(max_output_files={self.max_output_files})."
            )

    def _prepare_email_threads(
        self,
        email_files: List[str],
        warnings: Optional[List[Dict]] = None,
    ) -> Tuple[Dict[str, List[Dict]], Dict[str, int]]:
        """Extract and group email files into conversation threads."""
        email_data = []
        stats = {
            "parsed_total": 0,
            "failed_total": 0,
            "attachment_refs_total": 0,
        }
        
        for email_file in email_files:
            if email_file.lower().endswith('.msg'):
                data = self.email_extractor.extract_msg(email_file)
            else:
                data = self.email_extractor.extract_eml(email_file)
            
            if data:
                data['file_path'] = email_file
                data['attachments'] = data.get('attachments', []) or []
                email_data.append(data)
                stats["parsed_total"] += 1
                stats["attachment_refs_total"] += len(data['attachments'])
            else:
                _record_warning(
                    warnings,
                    'email_extract_failed',
                    'Could not parse email file; skipping file',
                    file=email_file,
                )
                stats["failed_total"] += 1
        
        # Group into threads
        return self.email_threader.group_emails(email_data), stats

    def _render_email_entry(
        self,
        email: Dict[str, Any],
        index: int,
        total: int,
    ) -> str:
        lines = [
            f"EMAIL {index} of {total}",
            "=" * 80,
            f"Subject: {email.get('subject', '')}",
            f"From: {email.get('from', '')}",
            f"To: {email.get('to', '')}",
            f"CC: {email.get('cc', '')}",
            f"Date: {email.get('date', '')}",
            f"Source: {os.path.basename(email.get('file_path', ''))}",
            "-" * 80,
            "",
            email.get('body', '') or "",
            "",
        ]

        attachments = email.get("attachments", []) or []
        if self.email_include_attachment_index:
            lines.append("ATTACHMENTS:")
            if attachments:
                for attachment in attachments:
                    size_bytes = attachment.get("size_bytes")
                    size_display = str(size_bytes) if size_bytes is not None else "unknown"
                    lines.append(
                        f"- {attachment.get('filename', 'unnamed_attachment')} "
                        f"(type={attachment.get('content_type', '')}, bytes={size_display})"
                    )
            else:
                lines.append("- none")
            lines.append("")

        lines.append("=" * 80)
        lines.append("")
        return "\n".join(lines)

    def _render_thread_block(
        self,
        thread_num: int,
        thread_key: str,
        emails: List[Dict[str, Any]],
    ) -> str:
        normalized_key = thread_key or "(no subject)"
        block_lines = [
            f"EMAIL THREAD {thread_num}",
            f"THREAD KEY: {normalized_key}",
            f"TOTAL EMAILS: {len(emails)}",
            "=" * 80,
            "",
        ]
        for idx, email in enumerate(emails, 1):
            try:
                block_lines.append(self._render_email_entry(email, idx, len(emails)))
            except Exception as exc:
                block_lines.append(
                    f"--- Email {idx}/{len(emails)}: RENDER FAILED ({exc}) ---\n"
                )
        return "\n".join(block_lines)

    def _write_email_outputs(
        self,
        threads: Dict[str, List[Dict]],
        output_path: str,
        group_name: str,
        current_output_count: int,
        warnings: Optional[List[Dict]] = None,
        run_logger: Optional[RunLogger] = None,
    ) -> Tuple[List[str], Dict[str, Any]]:
        if self.email_output_mode == "threaded":
            self._ensure_output_capacity(
                len(threads),
                current_output_count,
                f"group '{group_name}' email threads",
            )
            outputs = self._write_email_threads(threads, output_path, group_name)
            output_total_bytes = 0
            for file_path in outputs:
                try:
                    output_total_bytes += os.path.getsize(file_path)
                except OSError:
                    pass
            batch_map = {}
            for index, (thread_key, emails) in enumerate(sorted(threads.items()), 1):
                key = os.path.join(output_path, f"{group_name}_emails_thread{index}.txt")
                batch_map[key] = [{"thread_key": thread_key, "email_count": len(emails)}]
            return outputs, {
                "threads_total": len(threads),
                "batches_total": len(outputs),
                "output_total_bytes": output_total_bytes,
                "batch_to_threads": batch_map,
            }
        return self._write_email_batches(
            threads=threads,
            output_path=output_path,
            group_name=group_name,
            current_output_count=current_output_count,
            warnings=warnings,
            run_logger=run_logger,
        )

    def _write_email_threads(
        self,
        threads: Dict[str, List[Dict]],
        output_path: str,
        group_name: str,
    ) -> List[str]:
        """Write grouped email threads to text files."""
        os.makedirs(output_path, exist_ok=True)

        # Write thread files
        output_files = []
        for thread_num, (thread_key, emails) in enumerate(sorted(threads.items()), 1):
            output_file = os.path.join(output_path, f"{group_name}_emails_thread{thread_num}.txt")
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(f"GROUP: {group_name}\n")
                f.write(self._render_thread_block(thread_num, thread_key, emails))
            
            output_files.append(output_file)
            print(f"    Created: {os.path.basename(output_file)} ({len(emails)} emails)")
        
        return output_files

    def _write_email_batches(
        self,
        threads: Dict[str, List[Dict]],
        output_path: str,
        group_name: str,
        current_output_count: int,
        warnings: Optional[List[Dict]] = None,
        run_logger: Optional[RunLogger] = None,
    ) -> Tuple[List[str], Dict[str, Any]]:
        os.makedirs(output_path, exist_ok=True)
        max_batch_bytes = self.email_max_output_file_mb * 1024 * 1024
        max_batch_words = 50000  # netdoc word limit
        thread_blocks = []
        for thread_num, (thread_key, emails) in enumerate(sorted(threads.items()), 1):
            block_text = self._render_thread_block(thread_num, thread_key, emails)
            block_bytes = len(block_text.encode("utf-8"))
            block_words = len(block_text.split())
            thread_blocks.append(
                {
                    "thread_key": thread_key,
                    "email_count": len(emails),
                    "thread_num": thread_num,
                    "text": block_text,
                    "bytes": block_bytes,
                    "words": block_words,
                }
            )

        planned_batches: List[List[Dict[str, Any]]] = []
        current_batch: List[Dict[str, Any]] = []
        current_size = 0
        current_words = 0
        for block in thread_blocks:
            block_size = block["bytes"]
            block_words = block["words"]
            if current_batch and (current_size + block_size > max_batch_bytes or current_words + block_words > max_batch_words):
                planned_batches.append(current_batch)
                current_batch = []
                current_size = 0
                current_words = 0
            current_batch.append(block)
            current_size += block_size
            current_words += block_words
            if block_size > max_batch_bytes:
                _record_warning(
                    warnings,
                    "email_thread_exceeds_batch_cap",
                    "Email thread exceeds configured batch size cap; writing dedicated batch file",
                    thread_key=block["thread_key"],
                    group=group_name,
                    thread_bytes=block_size,
                    batch_limit_bytes=max_batch_bytes,
                )
            if block_words > max_batch_words:
                _record_warning(
                    warnings,
                    "email_thread_exceeds_word_cap",
                    "Email thread exceeds netdoc word limit (50000); writing dedicated batch file",
                    thread_key=block["thread_key"],
                    group=group_name,
                    thread_words=block_words,
                    batch_limit_words=max_batch_words,
                )
        if current_batch:
            planned_batches.append(current_batch)

        self._ensure_output_capacity(
            len(planned_batches),
            current_output_count,
            f"group '{group_name}' email batches",
        )

        output_files: List[str] = []
        output_total_bytes = 0
        batch_to_threads: Dict[str, List[Dict[str, Any]]] = {}
        for batch_num, batch_blocks in enumerate(planned_batches, 1):
            output_file = os.path.join(output_path, f"{group_name}_{self.email_batch_name_prefix}{batch_num}.txt")
            batch_word_count = sum(block["words"] for block in batch_blocks)
            with open(output_file, "w", encoding="utf-8") as handle:
                handle.write(f"EMAIL BATCH {batch_num}\n")
                handle.write(f"GROUP: {group_name}\n")
                handle.write(f"BATCH THREADS: {len(batch_blocks)}\n")
                handle.write(f"BATCH WORDS: {batch_word_count}\n")
                handle.write("=" * 80 + "\n\n")
                for idx, block in enumerate(batch_blocks):
                    if idx > 0:
                        handle.write("\n")
                    handle.write(block["text"])

            try:
                file_size = os.path.getsize(output_file)
            except OSError:
                file_size = 0
            output_total_bytes += file_size
            output_files.append(output_file)
            batch_to_threads[output_file] = [
                {
                    "thread_key": block["thread_key"],
                    "email_count": block["email_count"],
                }
                for block in batch_blocks
            ]
            if run_logger:
                run_logger.log(
                    "info",
                    "email_batch_written",
                    "Wrote email batch output",
                    group=group_name,
                    batch_file=output_file,
                    thread_count=len(batch_blocks),
                    bytes=file_size,
                    words=batch_word_count,
                )
            print(f"    Created: {os.path.basename(output_file)} ({len(batch_blocks)} threads, {batch_word_count} words)")

        return output_files, {
            "threads_total": len(threads),
            "batches_total": len(output_files),
            "output_total_bytes": output_total_bytes,
            "batch_to_threads": batch_to_threads,
        }

    def _process_emails(self, email_files: List[str], output_path: str, group_name: str) -> List[str]:
        """Backward-compatible wrapper for email processing."""
        email_threads, _ = self._prepare_email_threads(email_files)
        return self._write_email_threads(email_threads, output_path, group_name)
